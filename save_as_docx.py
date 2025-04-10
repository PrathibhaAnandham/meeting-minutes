from docx import Document

def save_as_docx(data, filename):
    doc = Document()
    for key, value in data.items():
        doc.add_heading(key, level=1)
        doc.add_paragraph(value)
    doc.save(filename)
